import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

# Process Mẫu 11
for name in ['11_Mau_Giay_trieu_tap.docx', '11_Mau_Giay_trieu_tap_PGD.docx']:
    template_path = os.path.join('webapp/templates/TOTRINH', name)
    doc = docx.Document(template_path)
    
    for p in doc.paragraphs:
        if 'Thời gian:' in p.text and 'Hội nghị' in p.text:
            p.text = ' - Thời gian: (TIME)'
        elif 'Địa điểm:' in p.text:
            p.text = ' - Địa điểm: (LOCATION)'
        elif 'a. Khối cơ quan' in p.text:
            p.text = ' (PARTICIPANTS)'
        elif 'b. Các đơn vị trực thuộc' in p.text:
            p.text = ''
        elif 'Các đơn vị tham dự đủ' in p.text:
            p.text = ' (ORGANIZATION)'
        elif 'Ban tổ chức Hội nghị đón tiếp' in p.text:
            p.text = ''
        elif 'Trong quá trình triển khai' in p.text:
            p.text = ''
            
    doc.save(template_path)
    print(f"Updated {name} successfully!")

# Process Mẫu 13
for name in ['13_Mau_Giay_moi.docx', '13_Mau_Giay_moi_PGD.docx']:
    template_path = os.path.join('webapp/templates/TOTRINH', name)
    doc = docx.Document(template_path)

    for p in doc.paragraphs:
        if 'Chủ trì:' in p.text:
            p.text = 'Chủ trì: (CHUTRI)'
        elif 'Thời gian:' in p.text:
            p.text = 'Thời gian: (TIME)'

    doc.save(template_path)
    print(f"Updated {name} successfully!")
