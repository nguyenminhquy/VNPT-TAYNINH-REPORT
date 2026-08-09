import os
import glob
import docx

def fix_title_in_paragraphs(paragraphs):
    modified = False
    for p in paragraphs:
        text = p.text
        if '(3)' in text:
            # We want to replace the whole string with the prefix + (TITLE)
            # Find the prefix
            if 'Về việc' in text:
                p.text = 'Về việc (TITLE)'
            elif 'Về' in text:
                p.text = 'Về (TITLE)'
            elif 'TỜ TRÌNH' in text:
                # sometimes it's on a new line
                p.text = p.text.replace(text, 'TỜ TRÌNH (TITLE)')
            else:
                p.text = '(TITLE)'
            modified = True
    return modified

template_dir = 'webapp/templates/TOTRINH'
for filepath in glob.glob(os.path.join(template_dir, '*.docx')):
    try:
        doc = docx.Document(filepath)
        mod1 = fix_title_in_paragraphs(doc.paragraphs)
        
        mod2 = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if fix_title_in_paragraphs(cell.paragraphs):
                        mod2 = True
                        
        if mod1 or mod2:
            doc.save(filepath)
            print(f"Fixed title in {os.path.basename(filepath)}")
    except Exception as e:
        print(f"Error in {filepath}: {e}")
