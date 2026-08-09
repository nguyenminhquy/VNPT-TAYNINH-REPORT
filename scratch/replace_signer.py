import docx
import os
import sys
import glob

sys.stdout.reconfigure(encoding='utf-8')

template_dir = 'webapp/templates/TOTRINH'
files = glob.glob(os.path.join(template_dir, '*.docx'))

for filepath in files:
    try:
        doc = docx.Document(filepath)
        modified = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'Họ và tên' in paragraph.text or 'Họ Và Tên' in paragraph.text:
                            paragraph.text = '(SIGNER)'
                            modified = True
        if modified:
            doc.save(filepath)
            print(f"Replaced signer in {os.path.basename(filepath)}")
    except Exception as e:
        print(f"Error processing {os.path.basename(filepath)}: {e}")
