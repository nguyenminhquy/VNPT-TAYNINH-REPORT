import copy
import json
from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.document import Document as DocumentType
from openpyxl import load_workbook

# Re‑use helper functions from existing monthly generator
from generate_monthly_report import (
    find_table_by_tag,
    write_table_matrix,
    raw_matrix,
)

# ---------------------------------------------------------------------------
# Configuration – template discovery
# ---------------------------------------------------------------------------
_CUR_DIR = Path(__file__).resolve().parent
try:
    _ROOT = _CUR_DIR.resolve().parents[1]
except IndexError:
    _ROOT = _CUR_DIR

_WEEKLY_TEMPLATE_CANDIDATES = [
    _CUR_DIR / "templates" / "MAT_LIEN_LAC" / "TTHT_BC_KTHT_MLL.docx",
    _ROOT / "templates" / "MAT_LIEN_LAC" / "TTHT_BC_KTHT_MLL.docx",
]
WEEKLY_TEMPLATE = next((p for p in _WEEKLY_TEMPLATE_CANDIDATES if p.is_file()), _WEEKLY_TEMPLATE_CANDIDATES[0])

EXPORT_DIR = _ROOT / "exports"
EXPORT_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Helper – download Excel blobs synchronously
# ---------------------------------------------------------------------------
import urllib.request
from datetime import datetime

# ---------------------------------------------------------------------------
# Core – load sources
# ---------------------------------------------------------------------------
def load_weekly_sources(blob_urls: Dict[str, str]) -> Dict[str, Any]:
    """Download two Excel files (keys ``weekly1`` and ``weekly2``) into ``DATA_DIR``
    and load them with ``openpyxl``.
    """
    data_dir = _ROOT / "data sample"
    data_dir.mkdir(parents=True, exist_ok=True)

    sources: Dict[str, Any] = {}
    for key, filename in {"weekly1": "weekly1.xlsx", "weekly2": "weekly2.xlsx"}.items():
        url = blob_urls.get(key)
        if not url:
            raise ValueError(f"Missing Blob URL for {key}")
        dest = data_dir / filename
        # Download synchronously
        urllib.request.urlretrieve(url, str(dest))
        sources[key] = load_workbook(dest, data_only=True, read_only=True)
    return sources

def extract_table_by_tag(sheet: Any, tag: str) -> list[list[Any]]:
    if getattr(sheet, "max_row", None) is None:
        # Prevent TypeError if sheet is empty or not properly loaded
        return []
        
    start_r, start_c = -1, -1
    for r in range(1, sheet.max_row + 1):
        for c in range(1, sheet.max_column + 1):
            val = sheet.cell(row=r, column=c).value
            if val and isinstance(val, str) and tag in val:
                start_r = r
                start_c = c
                break
        if start_r != -1:
            break
            
    if start_r == -1:
        return []
        
    max_c = start_c
    while max_c <= sheet.max_column and sheet.cell(row=start_r, column=max_c).value is not None:
        max_c += 1
    max_c -= 1
    
    # In some templates, tag is on a single cell row before the table.
    # In others, tag is in the header row.
    # We will extract starting from the row with the tag, going down.
    # If the user has empty rows, we should stop at the first completely empty row.
    max_r = start_r
    while max_r <= sheet.max_row:
        # Check if the whole row is empty from start_c to max_c
        is_empty = True
        for c in range(start_c, max_c + 1):
            if sheet.cell(row=max_r, column=c).value is not None:
                is_empty = False
                break
        if is_empty and max_r > start_r:
            break
        max_r += 1
    max_r -= 1
    
    matrix = []
    # Start extracting from the row below the tag if the tag row only contains the tag
    # If the tag is part of the header (i.e. other columns in that row have values), include it.
    tag_row_has_other_values = False
    for c in range(start_c + 1, max_c + 1):
        if sheet.cell(row=start_r, column=c).value is not None:
            tag_row_has_other_values = True
            break
            
    start_extract_row = start_r if tag_row_has_other_values else start_r + 1
    if start_extract_row > max_r:
        return []
        
    for r in range(start_extract_row, max_r + 1):
        row_data = []
        for c in range(start_c, max_c + 1):
            val = sheet.cell(row=r, column=c).value
            row_data.append("" if val is None else val)
        matrix.append(row_data)
        
    return matrix


def update_weekly_report(document: DocumentType, sources: Dict[str, Any]) -> None:
    """Populate tables in the weekly template by matching tags."""
    import re
    
    text = chr(10).join([p.text for p in document.paragraphs] + [c.text for t in document.tables for r in t.rows for c in r.cells])
    tags = list(set(re.findall(r'\(B\d.*?\)', text)))
    
    sheet1 = sources["weekly1"].active
    sheet2 = sources["weekly2"].active
    
    for tag in tags:
        tbl = find_table_by_tag(document, tag)
        if tbl is None:
            continue
            
        matrix = extract_table_by_tag(sheet1, tag)
        if not matrix:
            matrix = extract_table_by_tag(sheet2, tag)
            
        if matrix:
            write_table_matrix(tbl, matrix, start_row=0)

# ---------------------------------------------------------------------------
# Public API – generate the report and return the Path
# ---------------------------------------------------------------------------
def export_weekly(blob_urls: Dict[str, str]) -> Path:
    """Create the weekly loss‑of‑contact Word report.

    Returns the absolute path of the generated ``.docx`` file.
    """
    document = Document(str(WEEKLY_TEMPLATE))
    sources = load_weekly_sources(blob_urls)
    update_weekly_report(document, sources)

    out_path = EXPORT_DIR / f"Mất_liên_lạc_Tuần_{datetime.now():%Y%m%d}.docx"
    document.save(str(out_path))
    return out_path
