from docx import Document
from pathlib import Path
from pydantic import BaseModel
from typing import Dict, List
import io
import re

class DemoExportRequest(BaseModel):
    tables: Dict[str, List[List[str]]]

def remove_tags(doc, tag):
    """Xóa thẻ nhúng ra khỏi file Word"""
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            paragraph.text = paragraph.text.replace(tag, "")

def generate_demo_report_from_json(request: DemoExportRequest, template_path: str) -> io.BytesIO:
    doc = Document(template_path)
    
    # Process B1_TAM (MBB)
    if "B1_TAM" in request.tables and request.tables["B1_TAM"]:
        b1_data = request.tables["B1_TAM"]
        for i, paragraph in enumerate(doc.paragraphs):
            if "(B1_TAM)" in paragraph.text:
                table = doc.add_table(rows=1, cols=len(b1_data[0]))
                table.style = 'Table Grid'
                # Headers
                hdr_cells = table.rows[0].cells
                for col_idx, cell_value in enumerate(b1_data[0]):
                    hdr_cells[col_idx].text = str(cell_value)
                # Data
                for row_data in b1_data[1:]:
                    row_cells = table.add_row().cells
                    for col_idx, cell_value in enumerate(row_data):
                        row_cells[col_idx].text = str(cell_value)
                
                # Move table to after paragraph
                p_element = paragraph._p
                p_element.addnext(table._tbl)
                break
        remove_tags(doc, "(B1_TAM)")

    # Process B2_TAM (FBB)
    if "B2_TAM" in request.tables and request.tables["B2_TAM"]:
        b2_data = request.tables["B2_TAM"]
        for i, paragraph in enumerate(doc.paragraphs):
            if "(B2_TAM)" in paragraph.text:
                table = doc.add_table(rows=1, cols=len(b2_data[0]))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for col_idx, cell_value in enumerate(b2_data[0]):
                    hdr_cells[col_idx].text = str(cell_value)
                for row_data in b2_data[1:]:
                    row_cells = table.add_row().cells
                    for col_idx, cell_value in enumerate(row_data):
                        row_cells[col_idx].text = str(cell_value)
                        
                p_element = paragraph._p
                p_element.addnext(table._tbl)
                break
        remove_tags(doc, "(B2_TAM)")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
