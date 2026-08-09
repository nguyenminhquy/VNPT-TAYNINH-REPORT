from __future__ import annotations

import argparse
import copy
import json
import math
import re
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell


try:
    ROOT = Path(__file__).resolve().parents[1]
except IndexError:
    ROOT = Path(__file__).resolve().parent

DATA_DIR = ROOT / "data sample"
# Ưu tiên template.docx chuẩn trong /app/templates (Docker) hoặc root/templates, fallback sang file tên cũ
_CUR_DIR = Path(__file__).resolve().parent
_TEMPLATE_CANDIDATES = [
    _CUR_DIR / "templates" / "TTHT-BÁO CÁO THÁNG 7_UPDATE.docx",
    ROOT / "templates" / "TTHT-BÁO CÁO THÁNG 7_UPDATE.docx",
    ROOT / "webapp" / "templates" / "TTHT-BÁO CÁO THÁNG 7_UPDATE.docx",
]
TEMPLATE_FILE = next((p for p in _TEMPLATE_CANDIDATES if p.is_file()), _TEMPLATE_CANDIDATES[0])
EXPORT_DIR = ROOT / "exports"

FILES = {
    "mbb": "1. BÁO CÁO MBB_HUNG.xlsx",
    "fbb": "2. BÁO CÁO FBB_BAO.xlsx",
    "mytv": "3. BÁO CÁO MYTV_TÂN.xlsx",
    "ispeed": "5. BÁO CÁO ISPEED_QUOC.xlsx",
    "5s": "6. BÁO CÁO 5S NHÀ TRẠM_TÂN.xlsx",
    "xlsc": "7. BÁO CÁO XLSC_TUAN.xlsx",
    "appendix": "8. PHỤ LỤC 1_ HÂN.xlsx",
    "omc_tam": "9.HIỆN TRẠNG THIẾT BỊ_TÂM.xlsx",
    "omc_nhi": "10. BÁO CÁO BSC_NHI.xlsx",
    "cauhinh_quy": "11. KẾT QUẢ CẤU HÌNH TỰ ĐỘNG.xlsx",
    "ngoaivi_bao": "12. BÁO CÁO THIẾT BỊ NGOẠI VI.xlsx",
    "phutro_quy": "13. BÁO CÁO THIẾT BỊ PHỤ TRỢ.xlsx"
}


def clean(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y %H:%M:%S")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, bool):
        return "Có" if value else "Không"
    return str(value).strip()


def trim_number(value: float, decimals: int = 5) -> str:
    if math.isclose(value, round(value), abs_tol=1e-10):
        return str(int(round(value)))
    return f"{value:.{decimals}f}".rstrip("0").rstrip(".")


def format_cell(cell: Cell) -> str:
    value = cell.value
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return clean(value)
    if not isinstance(value, (int, float)) or isinstance(value, bool):
        return clean(value)

    number_format = str(cell.number_format or "General").split(";")[0]
    if "%" in number_format:
        decimals_match = re.search(r"\.([0#]+)%", number_format)
        decimals = len(decimals_match.group(1)) if decimals_match else 0
        return f"{value * 100:.{decimals}f}%"

    decimal_match = re.search(r"\.([0#]+)", number_format)
    decimals = len(decimal_match.group(1)) if decimal_match else None
    use_grouping = "," in number_format.split(".")[0]
    if decimals is not None:
        return f"{value:,.{decimals}f}" if use_grouping else f"{value:.{decimals}f}"
    if use_grouping:
        return f"{value:,.0f}"
    return trim_number(float(value))


def worksheet_matrix(
    sheet: Any,
    min_row: int,
    max_row: int,
    min_col: int,
    max_col: int,
) -> list[list[str]]:
    return [
        [format_cell(sheet.cell(row=row, column=column)) for column in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def raw_matrix(
    sheet: Any,
    min_row: int,
    max_row: int,
    min_col: int,
    max_col: int,
) -> list[list[Any]]:
    return [
        [sheet.cell(row=row, column=column).value for column in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def load_sources() -> dict[str, Any]:
    sources: dict[str, Any] = {}
    for key, filename in FILES.items():
        if (DATA_DIR / filename).is_file():
            sources[key] = load_workbook(DATA_DIR / filename, data_only=True, read_only=True)
    return sources


def first_run_properties(paragraph: Paragraph) -> Any | None:
    for run in paragraph._p.xpath(".//w:r"):
        run_properties = run.find(qn("w:rPr"))
        if run_properties is not None:
            return copy.deepcopy(run_properties)
    return None


def append_text_run(paragraph: Paragraph, value: str, run_properties: Any | None) -> None:
    lines = clean(value).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    run = OxmlElement("w:r")
    if run_properties is not None:
        run.append(copy.deepcopy(run_properties))

    for index, line in enumerate(lines):
        if index:
            run.append(OxmlElement("w:br"))
        text = OxmlElement("w:t")
        if line.startswith(" ") or line.endswith(" "):
            text.set(qn("xml:space"), "preserve")
        text.text = line
        run.append(text)
    paragraph._p.append(run)


def replace_paragraph(paragraph: Paragraph, value: Any) -> None:
    run_properties = first_run_properties(paragraph)
    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)
    append_text_run(paragraph, clean(value), run_properties)


def replace_cell(cell: _Cell, value: Any) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        paragraph_element = OxmlElement("w:p")
        cell._tc.append(paragraph_element)
        paragraphs = [Paragraph(paragraph_element, cell)]

    first = paragraphs[0]
    replace_paragraph(first, value)
    for paragraph in paragraphs[1:]:
        cell._tc.remove(paragraph._p)


def write_table_matrix(
    table: Table,
    matrix: Sequence[Sequence[Any]],
    start_row: int = 0,
    start_col: int = 0,
) -> None:
    touched: set[Any] = set()
    for source_row_index, source_row in enumerate(matrix):
        target_row_index = start_row + source_row_index
        if target_row_index >= len(table.rows):
            break
        cells = table.rows[target_row_index].cells
        for source_col_index, value in enumerate(source_row):
            target_col_index = start_col + source_col_index
            if target_col_index >= len(cells):
                break
            cell = cells[target_col_index]
            identity = cell._tc
            if identity in touched:
                continue
            replace_cell(cell, value)
            touched.add(identity)


def resize_table_rows(table: Table, desired_rows: int) -> None:
    while len(table.rows) < desired_rows:
        source_row = table.rows[-1]._tr
        table._tbl.append(copy.deepcopy(source_row))
    while len(table.rows) > desired_rows:
        table._tbl.remove(table.rows[-1]._tr)


def as_number(value: Any) -> float:
    if value is None or value == "":
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace("%", "").replace(",", "").strip()
    try:
        return float(text)
    except ValueError:
        return 0.0


def percent(value: Any, decimals: int = 2, stored_as_percent: bool = False) -> str:
    number = as_number(value)
    if not stored_as_percent:
        number *= 100
    return f"{number:.{decimals}f}%"


def integer(value: Any, dash_zero: bool = False) -> str:
    number = as_number(value)
    if dash_zero and math.isclose(number, 0.0, abs_tol=1e-9):
        return "-"
    return f"{round(number):,}"


def decimal(value: Any, decimals: int = 2) -> str:
    return f"{as_number(value):.{decimals}f}"


def evaluate_target(value: Any, target: float = 99.0) -> str:
    return "Đạt" if as_number(value) >= target else "Không đạt"



def find_table_by_tag(document: DocumentType, tag: str, offset: int = 0) -> Any:
    table_idx = 0
    for element in document.element.body:
        if element.tag.endswith('p'):
            if element.text and tag in element.text:
                return document.tables[table_idx + offset]
        elif element.tag.endswith('tbl'):
            table = document.tables[table_idx]
            text = chr(10).join([c.text for r in table.rows for c in r.cells])
            if tag in text:
                return document.tables[table_idx + offset]
            table_idx += 1
    return None

def extract_dynamic_table(document: DocumentType, sheet: Any, tag: str, anchor_text: str = None, start_col_idx: int = 1, apply_formatting=None, word_start_row=None, row_offset=0) -> None:
    t = find_table_by_tag(document, tag)
    if not t: return
    
    num_cols = len(t.rows[0].cells)
    
    start_row = None
    start_col = None
    for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
        for col_idx, cell in enumerate(row, 1):
            if isinstance(cell, str):
                if anchor_text and (anchor_text in cell or anchor_text.replace('â','a') in cell.replace('â','a')):
                    start_row = row_idx
                    start_col = start_col_idx
                    break
                elif not anchor_text and tag in cell:
                    start_row = row_idx
                    start_col = start_col_idx
                    break
        if start_row: break
    
    if not start_row: return
    
    start_row += row_offset
        
    matrix = []
    for row_idx, row in enumerate(sheet.iter_rows(min_row=start_row, max_col=start_col + num_cols - 1, values_only=True), start_row):
        if all(c is None or str(c).strip() == "" for c in row[start_col - 1 : start_col - 1 + num_cols]):
            break
        row_data = list(row[start_col - 1 : start_col - 1 + num_cols])
        if apply_formatting:
            row_data = apply_formatting(row_data, row_idx - start_row)
        matrix.append(row_data)
        
    if word_start_row is None:
        word_start_row = 1
        for i in range(len(t.rows)):
            row_text = " ".join(c.text for c in t.rows[i].cells)
            if anchor_text and (anchor_text in row_text or anchor_text.replace('â','a') in row_text.replace('â','a')):
                word_start_row = i
                break
                
    try:
        write_table_matrix(t, matrix, start_row=word_start_row)
        print(f"Successfully wrote {tag} dynamically")
    except Exception as e:
        print(f"Error {tag}:", e)

def remove_tags(document: DocumentType) -> None:
    import re
    pattern = r'\([Bb]\d+G?(?:_[^\)]+)?\)'
    for paragraph in document.paragraphs:
        if re.search(pattern, paragraph.text):
            text = re.sub(pattern, '', paragraph.text)
            replace_paragraph(paragraph, text.strip())
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if re.search(pattern, paragraph.text):
                        text = re.sub(pattern, '', paragraph.text)
                        replace_paragraph(paragraph, text.strip())

def update_mbb_fbb_mytv(document: DocumentType, sources: dict[str, Any]) -> None:
    mbb = sources.get("mbb")
    fbb = sources.get("fbb")
    mytv = sources.get("mytv")

    if mbb:
        try:
            def fmt_b15(row, idx):
                if idx > 0:
                    row[2] = decimal(row[2]) if len(row) > 2 else row[2]
                    row[3] = decimal(row[3]) if len(row) > 3 else row[3]
                return row
            extract_dynamic_table(document, mbb["Kết quả chung"], "B15_HUNG", "Toàn quốc", apply_formatting=fmt_b15, word_start_row=1)
        except Exception as e: print("mbb 15:", e)

        try:
            extract_dynamic_table(document, mbb["So sánh các tỉnh"], "B16_HUNG", "Thanh Hóa", word_start_row=1)
        except Exception as e: print("mbb 16:", e)

        try:
            extract_dynamic_table(document, mbb["Kết quả chi tiết"], "B17_HUNG", "MBB QoS", word_start_row=1)
        except Exception as e: print("mbb 17:", e)

    if fbb:
        try:
            extract_dynamic_table(document, fbb["Thông tin chung"], "B18_BAO", "FBB QoS", word_start_row=1)
            extract_dynamic_table(document, fbb["Thông tin chung"], "B19_BAO", "FBB QoE", word_start_row=1)
        except Exception as e: print("fbb:", e)

    if mytv:
        try:
            def fmt_mytv(row, idx):
                if len(row) >= 8:
                    total = row[6]
                    return [clean(row[0]), clean(row[1]), clean(row[3]), clean(row[4]), clean(row[5]), clean(total), evaluate_target(total), clean(row[7])]
                return row
            extract_dynamic_table(document, mytv["Sheet1"], "B14", "MyTV QoS", start_col_idx=1, apply_formatting=fmt_mytv, word_start_row=1)
            extract_dynamic_table(document, mytv["Sheet1"], "B21_TAN", "MyTV QoS", start_col_idx=1, apply_formatting=fmt_mytv, word_start_row=1)
        except Exception as e: print("mytv:", e)

def mll_table_matrix(sheet: Any) -> tuple[list[list[str]], dict[str, Any]]:
    raw = raw_matrix(sheet, 2, 12, 1, 18)
    title = clean(raw[0][0])
    matrix: list[list[str]] = [[title]]

    source_columns = [0, 1, 2, 3, 4, 5, 7, 8, 9, 11, 12, 13, 15, 16, 17]
    for row in raw[3:11]:
        target: list[str] = []
        for target_index, source_index in enumerate(source_columns):
            value = row[source_index]
            if target_index in {0, 1}:
                target.append(clean(value))
            elif target_index == 2:
                target.append(integer(value))
            elif 3 <= target_index <= 11:
                target.append(integer(value, dash_zero=True))
            elif target_index in {12, 13}:
                target.append(integer(value))
            else:
                target.append(decimal(value))
        matrix.append(target)

    overall = raw[3]
    teams = raw[4:11]
    metrics = {
        "title": title,
        "week": re.search(r"TUẦN\s+(\d+)", title, re.IGNORECASE),
        "total": as_number(overall[15]),
        "average": as_number(overall[17]),
        "teams": [(clean(row[1]), as_number(row[17])) for row in teams],
        "cause_power": as_number(overall[3]) + as_number(overall[7]) + as_number(overall[11]),
        "cause_equipment": as_number(overall[4]) + as_number(overall[8]) + as_number(overall[12]),
        "cause_transmission": as_number(overall[5]) + as_number(overall[9]) + as_number(overall[13]),
    }
    return matrix, metrics


def update_mll(document: DocumentType, sources: dict[str, Any]) -> str:
    return ""

def update_ispeed(document: DocumentType, sources: dict[str, Any]) -> None:
    if "ispeed" not in sources: return
    try:
        extract_dynamic_table(document, sources["ispeed"]["Báo cáo"], "B22_QUOC", "Tân Ninh", start_col_idx=2, word_start_row=1)
    except Exception as e:
        print("Error ispeed:", e)

def update_5s(document: DocumentType, sources: dict[str, Any]) -> None:
    if "5s" not in sources: return
    try:
        extract_dynamic_table(document, sources["5s"]["Sheet1"], "B23_TAN", "5S NHÀ TRẠM", start_col_idx=1, word_start_row=1, row_offset=2)
        extract_dynamic_table(document, sources["5s"]["Sheet1"], "B24_TAN", "MÁY LẠNH", start_col_idx=1, word_start_row=1, row_offset=2)
        extract_dynamic_table(document, sources["5s"]["Sheet1"], "B25_TAN", "AP/OTB", start_col_idx=1, word_start_row=1, row_offset=2)
    except Exception as e:
        print("Error 5s:", e)

def extract_by_tag_in_sheet(document, wb, tag, max_rows=20, max_cols=10):
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
            for col_idx, cell in enumerate(row, 1):
                if isinstance(cell, str) and (tag in cell or cell.replace("0", "O") == tag.replace("0", "O")):
                    # Tag found! The data is typically below or left of it.
                    # Since we know the column of the tag is the far right, we extract from col 2 to col_idx-1
                    # And rows from row_idx to row_idx+max_rows
                    start_col = 2
                    end_col = col_idx - 1
                    if end_col < start_col: end_col = col_idx
                    start_row = row_idx
                    # Find end_row by looking for empty rows
                    end_row = row_idx + max_rows
                    
                    matrix = raw_matrix(sheet, start_row, end_row, start_col, end_col)
                    # Filter out completely empty rows
                    matrix = [r for r in matrix if any(c != "" for c in r)]
                    
                    t = find_table_by_tag(document, tag)
                    if t: 
                        try:
                            write_table_matrix(t, matrix, start_row=1)
                            print(f"Successfully wrote {tag}")
                        except Exception as e:
                            print(f"Error writing {tag}: {e}")
                    return True
    return False

def update_xlsc(document: DocumentType, sources: dict[str, Any]) -> None:
    if "xlsc" in sources:
        try:
            wb = sources["xlsc"]
            for tag in ["B26_TUAN", "B27_TUAN", "B28_TUAN", "B29_TUAN", "B30_TUAN"]:
                extract_by_tag_in_sheet(document, wb, tag, max_rows=15)
        except Exception as e:
            print("Error xlsc:", e)

def update_others(document: DocumentType, sources: dict[str, Any]) -> None:
    if "phutro_quy" in sources:
        try:
            extract_by_tag_in_sheet(document, sources["phutro_quy"], "B31_QUY", max_rows=10)
        except Exception as e: print("Error phutro_quy:", e)
        
    if "cauhinh_quy" in sources:
        try:
            # Handle typo B13_QUY -> B32_QUY
            extract_by_tag_in_sheet(document, sources["cauhinh_quy"], "B13_QUY", max_rows=15)
            # Try B32_QUY as well
            extract_by_tag_in_sheet(document, sources["cauhinh_quy"], "B32_QUY", max_rows=15)
        except Exception as e: print("Error cauhinh_quy:", e)
        
    if "ngoaivi_bao" in sources:
        try:
            extract_by_tag_in_sheet(document, sources["ngoaivi_bao"], "B33_BAO", max_rows=15, max_cols=16)
            extract_by_tag_in_sheet(document, sources["ngoaivi_bao"], "B34_BAO", max_rows=15, max_cols=16)
        except Exception as e: print("Error ngoaivi_bao:", e)

def update_tam_nhi(document: DocumentType, sources: dict[str, Any]) -> None:
    if "omc_nhi" in sources:
        try:
            wb = sources["omc_nhi"]
            extract_dynamic_table(document, wb["BANG 1 "], "B8_NHI", anchor_text=None, start_col_idx=1, word_start_row=1)
            extract_dynamic_table(document, wb["BANG3"], "B9_NHI", anchor_text=None, start_col_idx=1, word_start_row=1)
        except Exception as e:
            print("Error processing omc_nhi:", e)
            
    if "omc_tam" in sources:
        try:
            wb = sources["omc_tam"]
            
            # B1_TAM has a complex header where Tây Ninh data is embedded in merged header cells (Row 0 and 1).
            # The actual data rows start at "Bến Lức" (Row 2 in Word).
            # The Word table has 10 columns, but it does NOT have an STT column (Tây Ninh is Col 1).
            # So start_col_idx=2 to skip the STT column in Excel.
            extract_dynamic_table(document, wb["01_MANE_CSG"], "B1_TAM", anchor_text="Bến Lức", start_col_idx=2, word_start_row=2)
            
            # B2_TAM to B7_TAM have normal data rows starting with "Tây Ninh" at Row 1.
            # And they DO have an STT column in Word. So start_col_idx=1.
            mapping = {
                "02_OLT": "B2_TAM",
                "03_L2SW": "B3_TAM",
                "05_3G": "B4_TAM",
                "06_4G": "B5_TAM",
                "07_5G": "B6_TAM",
                "08_DLU": "B7_TAM"
            }
            for sheet_name, tag in mapping.items():
                extract_dynamic_table(document, wb[sheet_name], tag, anchor_text="Tây Ninh", start_col_idx=1, word_start_row=1)
                
        except Exception as e:
            print("Error processing omc_tam:", e)

def update_appendix(document: DocumentType, sources: dict[str, Any]) -> None:
    wb = sources["appendix"]
    sheet = wb.worksheets[0] if hasattr(wb, "worksheets") and len(wb.worksheets) > 0 else wb["Báo Cáo Sự Cố Trạm"]

    def is_valid_stt(val: Any) -> bool:
        if val is None:
            return False
        s = str(val).strip()
        if not s:
            return False
        try:
            return float(s) > 0
        except ValueError:
            return False

    all_rows = list(sheet.iter_rows(values_only=False))
    data_indices = [
        i for i, row in enumerate(all_rows)
        if len(row) > 0 and is_valid_stt(row[0].value)
    ]
    if not data_indices:
        return

    min_idx = min(data_indices)
    max_idx = max(data_indices)

    col2_header = str(all_rows[min_idx - 1][1].value or "").strip().lower() if min_idx > 0 and len(all_rows[min_idx - 1]) > 1 else ""
    if "sự cố" in col2_header or "su co" in col2_header:
        target_cols = [0, 2, 3, 4, 5, 6, 7, 8, 9, 10]
    else:
        target_cols = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9]

    matrix = []
    for idx in range(min_idx, max_idx + 1):
        row_cells = all_rows[idx]
        matrix.append([
            format_cell(row_cells[c]) if c < len(row_cells) else ""
            for c in target_cols
        ])

    t = find_table_by_tag(document, "B32_HAN")
    if t:
        resize_table_rows(t, len(matrix))
        write_table_matrix(t, matrix)


def iter_story_roots(document: DocumentType) -> Iterable[Any]:
    yield document.element
    seen: set[int] = set()
    for section in document.sections:
        for story in (section.header, section.footer):
            identity = id(story.part)
            if identity not in seen:
                seen.add(identity)
                yield story._element


def flatten_simple_link_fields(root: Any) -> int:
    count = 0
    instruction_attribute = qn("w:instr")
    for field in list(root.iter(qn("w:fldSimple"))):
        if not clean(field.get(instruction_attribute)).upper().startswith("LINK "):
            continue
        parent = field.getparent()
        insert_at = parent.index(field)
        for child in list(field):
            field.remove(child)
            parent.insert(insert_at, child)
            insert_at += 1
        parent.remove(field)
        count += 1
    return count


def flatten_complex_link_fields(root: Any) -> int:
    contexts: list[dict[str, Any]] = []
    completed: list[dict[str, Any]] = []

    for run in list(root.iter(qn("w:r"))):
        for child in list(run):
            if child.tag == qn("w:fldChar"):
                marker = child.get(qn("w:fldCharType"))
                if marker == "begin":
                    contexts.append({"instruction": [], "code_runs": {run}, "phase": "code"})
                elif marker == "separate" and contexts:
                    contexts[-1]["code_runs"].add(run)
                    contexts[-1]["phase"] = "result"
                elif marker == "end" and contexts:
                    context = contexts.pop()
                    context["end_run"] = run
                    completed.append(context)
            elif child.tag == qn("w:instrText") and contexts:
                context = contexts[-1]
                if context["phase"] == "code":
                    context["instruction"].append(child.text or "")
                    context["code_runs"].add(run)

    # The retained template contains one unterminated 5S LINK field. Its code
    # can still be safely removed because the linked result table follows it.
    completed.extend(contexts)

    link_contexts = [
        context
        for context in completed
        if "".join(context["instruction"]).lstrip().upper().startswith("LINK ")
    ]
    runs_to_clean: set[Any] = set()
    for context in link_contexts:
        runs_to_clean.update(context["code_runs"])
        if context.get("end_run") is not None:
            runs_to_clean.add(context["end_run"])

    meaningful_tags = {
        qn("w:t"),
        qn("w:tab"),
        qn("w:br"),
        qn("w:drawing"),
        qn("w:object"),
        qn("w:pict"),
    }
    for run in runs_to_clean:
        for node in list(run.iter()):
            if node.tag not in {qn("w:fldChar"), qn("w:instrText")}:
                continue
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
        if not any(node.tag in meaningful_tags for node in run.iter()):
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)

    return len(link_contexts)


def flatten_link_fields(document: DocumentType) -> int:
    count = 0
    for root in iter_story_roots(document):
        count += flatten_simple_link_fields(root)
        count += flatten_complex_link_fields(root)
    return count


def replace_report_week(document: DocumentType, week: str) -> None:
    if not week:
        return
    plan_match = re.search(r"tuần\s+(\d+)", document.paragraphs[18].text, re.IGNORECASE)
    plan_week = plan_match.group(1) if plan_match else str(int(week) + 1)
    replacements = {
        1: f"V/v thực hiện công việc trọng tâm trong tuần {week} năm 2026",
        2: f"và kế hoạch thực hiện nhiệm vụ tuần {plan_week}",
        4: f"Trung tâm Hạ tầng báo cáo kết quả thực hiện công việc trọng tâm trong tuần {week} năm 2026 như sau:",
        116: f"Trên đây là báo cáo kết quả thực hiện công việc tuần {week} năm 2026.",
    }
    for index, value in replacements.items():
        replace_paragraph(document.paragraphs[index], value)


def generate(output: Path | None = None) -> dict[str, Any]:
    if not TEMPLATE_FILE.is_file():
        raise FileNotFoundError(f"Không tìm thấy mẫu Word: {TEMPLATE_FILE}")

    sources = load_sources()
    try:
        document = Document(TEMPLATE_FILE)
        update_mbb_fbb_mytv(document, sources)
        week = update_mll(document, sources)
        update_ispeed(document, sources)
        update_5s(document, sources)
        update_xlsc(document, sources)
        update_others(document, sources)
        update_tam_nhi(document, sources)
        remove_tags(document)
        replace_report_week(document, week)
        flattened_links = flatten_link_fields(document)

        EXPORT_DIR.mkdir(parents=True, exist_ok=True)
        if output is None:
            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            week_label = f" tuần {week}" if week else ""
            output = EXPORT_DIR / f"TTHT Báo cáo công việc{week_label} - cập nhật {timestamp}.docx"
        else:
            output = output.resolve()
            output.parent.mkdir(parents=True, exist_ok=True)

        document.save(output)
    finally:
        for workbook in sources.values():
            workbook.close()

    return {
        "path": str(output),
        "filename": output.name,
        "size": output.stat().st_size,
        "week": week,
        "flattenedLinks": flattened_links,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the weekly VNPT Word report")
    parser.add_argument("--output", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    print(json.dumps(generate(args.output), ensure_ascii=True))


if __name__ == "__main__":
    main()
