from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def clean_text(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def paragraph_record(paragraph: Paragraph, index: int) -> dict[str, Any]:
    runs = []
    for run in paragraph.runs:
        text = clean_text(run.text)
        if not text:
            continue
        runs.append(
            {
                "text": text,
                "bold": run.bold,
                "italic": run.italic,
                "font": run.font.name,
                "sizePt": run.font.size.pt if run.font.size else None,
            }
        )
    return {
        "index": index,
        "text": clean_text(paragraph.text),
        "style": paragraph.style.name if paragraph.style else "",
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "runs": runs,
    }


def table_record(table: Table, index: int) -> dict[str, Any]:
    rows = []
    for row_index, row in enumerate(table.rows):
        cells = []
        for cell_index, cell in enumerate(row.cells):
            text = clean_text("\n".join(paragraph.text for paragraph in cell.paragraphs))
            cells.append({"index": cell_index, "text": text})
        rows.append({"index": row_index, "cells": cells})
    return {
        "index": index,
        "rows": len(table.rows),
        "columns": len(table.columns),
        "style": table.style.name if table.style else "",
        "data": rows,
    }


def iter_body_blocks(document: Document):
    paragraph_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", paragraph_index, Paragraph(child, document)
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            yield "table", table_index, Table(child, document)
            table_index += 1


def main() -> None:
    source = Path(sys.argv[1]).resolve()
    output = Path(sys.argv[2]).resolve()
    document = Document(source)

    inventory: dict[str, Any] = {
        "source": str(source),
        "sections": len(document.sections),
        "body": [],
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": [],
    }

    for kind, index, block in iter_body_blocks(document):
        if kind == "paragraph":
            record = paragraph_record(block, index)
            inventory["paragraphs"].append(record)
            inventory["body"].append({"kind": kind, "index": index, "text": record["text"]})
        else:
            record = table_record(block, index)
            inventory["tables"].append(record)
            preview = " | ".join(cell["text"] for cell in record["data"][0]["cells"]) if record["data"] else ""
            inventory["body"].append({"kind": kind, "index": index, "preview": preview})

    for section_index, section in enumerate(document.sections):
        inventory["headers"].append(
            {
                "section": section_index,
                "paragraphs": [clean_text(item.text) for item in section.header.paragraphs],
            }
        )
        inventory["footers"].append(
            {
                "section": section_index,
                "paragraphs": [clean_text(item.text) for item in section.footer.paragraphs],
            }
        )

    output.write_text(json.dumps(inventory, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {
                "paragraphs": len(inventory["paragraphs"]),
                "tables": len(inventory["tables"]),
                "sections": inventory["sections"],
                "output": str(output),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
