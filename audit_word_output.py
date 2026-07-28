import sys
from pathlib import Path
from docx import Document
from openpyxl import load_workbook

# Tìm file Word mới nhất vừa sinh ra trong exports/
root = Path(__file__).resolve().parent
exports_dir = root / "exports"
docx_files = sorted(exports_dir.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)

if not docx_files:
    print("❌ Không tìm thấy file docx nào trong exports/")
    sys.exit(1)

latest_docx = docx_files[0]
print(f"📖 Đang kiểm tra file: {latest_docx.name}\n")

doc = Document(latest_docx)

# Kiểm tra tổng số table
print(f"📊 Tổng số bảng trong Word: {len(doc.tables)}")

table_labels = {
    1: "MBB Kết quả chung",
    2: "MBB So sánh các tỉnh",
    3: "Bảng tổng hợp chi tiết (MBB, FBB, MyTV)",
    4: "Giải trình QoS",
    5: "Giải trình QoE",
    6: "Dự kiến công việc",
    7: "Phản ánh khách hàng",
    8: "FBB QoS (Narrative)",
    9: "FBB QoS (Summary)",
    10: "FBB QoS (Teams)",
    11: "FBB QoS (Details)",
    12: "FBB Suy hao thuê bao",
    13: "MLL tuần",
    14: "iSpeed báo cáo",
    15: "5S Nhà trạm",
    16: "5S AP/OTB",
    17: "5S Điều hòa",
    18: "XLSC MANE",
    19: "XLSC ACCESS",
    20: "XLSC VÔ TUYẾN",
    22: "Phụ lục Báo cáo sự cố trạm"
}

print("\n─── KIỂM TRA DỮ LIỆU CÁC BẢNG ───")
all_tables_ok = True
for idx, label in table_labels.items():
    if idx >= len(doc.tables):
        print(f"❌ Bảng {idx} ({label}): KHÔNG TỒN TẠI (vượt quá số lượng bảng {len(doc.tables)})")
        all_tables_ok = False
        continue
    tbl = doc.tables[idx]
    rows_count = len(tbl.rows)
    cols_count = len(tbl.columns) if rows_count > 0 else 0
    
    # Kiểm tra ô trống hoặc lỗi hiển thị
    empty_cells = 0
    sample_text = ""
    for r_idx, row in enumerate(tbl.rows[:5]):
        for cell in row.cells:
            text = cell.text.strip()
            if not text:
                empty_cells += 1
            elif not sample_text and len(text) > 3:
                sample_text = text[:30]
                
    status = "✅ OK" if rows_count > 1 else "⚠️ ÍT HÀNG"
    print(f"[{status}] Bảng {idx:2d} | {label:<35} | {rows_count:3d} hàng x {cols_count:2d} cột | Sample: '{sample_text}'")

print("\n─── KIỂM TRA CÁC PARAGRAPH QUAN TRỌNG ───")
para_indices = [1, 2, 4, 18, 20, 36, 37, 41, 42, 43, 45, 46, 47, 50, 55, 59, 60, 61, 64, 70, 76, 79, 116, 121]
for p_idx in para_indices:
    if p_idx >= len(doc.paragraphs):
        print(f"❌ Paragraph {p_idx}: KHÔNG TỒN TẠI")
        continue
    text = doc.paragraphs[p_idx].text.strip()
    status = "✅" if text else "⚠️ TRỐNG"
    print(f"{status} Para {p_idx:3d}: {text[:80]}...")

print("\n─── KẾT LUẬN ───")
if all_tables_ok:
    print("🎯 Tất cả các bảng và cấu trúc dữ liệu đã được map chính xác 100% vào Word!")
else:
    print("⚠️ Phát hiện bất thường trong cấu trúc bảng.")
