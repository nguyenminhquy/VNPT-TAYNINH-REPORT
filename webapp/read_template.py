import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')
doc_path = sys.argv[1]
doc = docx.Document(doc_path)

print("---PARAGRAPHS---")
for i, p in enumerate(doc.paragraphs):
    print(f"P{i}: '{p.text}'")

print("---TABLES---")
for t_idx, t in enumerate(doc.tables):
    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            print(f"T{t_idx} R{r_idx} C{c_idx}:")
            for i, p in enumerate(cell.paragraphs):
                print(f"  P{i}: '{p.text}'")
