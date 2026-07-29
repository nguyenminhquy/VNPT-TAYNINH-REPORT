import os
import sys
import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def prepare():
    src = r'c:\Users\quyng\Documents\TOOL TỔNG HỢP BÁO CÁO\VNPT REPORT\templates\To_trinh.docx'
    dst = r'c:\Users\quyng\Documents\TOOL TỔNG HỢP BÁO CÁO\VNPT REPORT\webapp\templates\ToTrinh_Template.docx'
    
    doc = docx.Document(src)
    
    # Check if headers exist
    has_header = False
    for p in doc.paragraphs:
        if 'TẬP ĐOÀN' in p.text or 'CỘNG HÒA' in p.text:
            has_header = True
            break
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if 'TẬP ĐOÀN' in c.text or 'CỘNG HÒA' in c.text:
                    has_header = True
                    break
    
    # If no header, add it at the top
    if not has_header:
        # Add Header table
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Mm(70)
        table.columns[1].width = Mm(95)
        
        cell_left = table.cell(0, 0)
        p_left = cell_left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_left1 = p_left.add_run('TẬP ĐOÀN BƯU CHÍNH VIỄN THÔNG VIỆT NAM\n')
        r_left1.font.size = Pt(12)
        r_left2 = p_left.add_run('VNPT TÂY NINH\n')
        r_left2.bold = True
        r_left2.font.size = Pt(13)
        r_left3 = p_left.add_run('Số: {{SO_TO_TRINH}}')
        r_left3.font.size = Pt(13)
        
        cell_right = table.cell(0, 1)
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_right1 = p_right.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n')
        r_right1.bold = True
        r_right1.font.size = Pt(13)
        r_right2 = p_right.add_run('Độc lập - Tự do - Hạnh phúc\n')
        r_right2.bold = True
        r_right2.font.size = Pt(14)
        r_right2.underline = True
        r_right3 = p_right.add_run('{{NGAY_THANG}}')
        r_right3.font.size = Pt(13)
        
        # Move the table to the top (this is a bit hacky in python-docx, we'll insert before first paragraph)
        first_p = doc.paragraphs[0]
        p_element = first_p._p
        tbl_element = table._tbl
        p_element.addprevious(tbl_element)
    
    # Replace content inside tables or paragraphs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if 'Về việc' in p.text and '……' in p.text:
                        p.text = 'Về việc {{TIEU_DE}}'
                        p.style.font.bold = True
                    if 'Căn cứ tờ trình số 1937' in p.text:
                        p.text = '{{CAN_CU}}'
                    if 'Để tạo điều kiện thuận lợi' in p.text:
                        p.text = '{{NOI_DUNG}}'
                    if 'Tổ Khai thác Hệ thống kính đề nghị' in p.text:
                        p.text = '{{DE_XUAT}}'
                        
    for p in doc.paragraphs:
        if 'Về việc' in p.text and '……' in p.text:
            p.text = 'Về việc {{TIEU_DE}}'
            p.style.font.bold = True
        if 'Căn cứ tờ trình số 1937' in p.text:
            p.text = '{{CAN_CU}}'
        if 'Để tạo điều kiện thuận lợi' in p.text:
            p.text = '{{NOI_DUNG}}'
        if 'Tổ Khai thác Hệ thống kính đề nghị' in p.text:
            p.text = '{{DE_XUAT}}'

    # Check if footer exists
    has_footer = False
    for p in doc.paragraphs:
        if 'NGƯỜI LẬP' in p.text or 'NGƯỜI PHÊ DUYỆT' in p.text:
            has_footer = True
            break
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if 'NGƯỜI LẬP' in c.text or 'NGƯỜI PHÊ DUYỆT' in c.text:
                    has_footer = True
                    break
                    
    if not has_footer:
        doc.add_paragraph()
        table_footer = doc.add_table(rows=1, cols=3)
        table_footer.columns[0].width = Mm(50)
        table_footer.columns[1].width = Mm(50)
        table_footer.columns[2].width = Mm(65)
        
        cell_f_left = table_footer.cell(0, 0)
        p_fl1 = cell_f_left.paragraphs[0]
        r_fl1 = p_fl1.add_run('Nơi nhận:\n')
        r_fl1.bold = True
        p_fl1.add_run('{{NOI_NHAN}}')
        
        cell_f_mid = table_footer.cell(0, 1)
        p_fm1 = cell_f_mid.paragraphs[0]
        p_fm1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_fm1 = p_fm1.add_run('NGƯỜI LẬP\n\n\n\n\n')
        r_fm1.bold = True
        p_fm1.add_run('{{NGUOI_LAP}}').bold = True
        
        cell_f_right = table_footer.cell(0, 2)
        p_fr1 = cell_f_right.paragraphs[0]
        p_fr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_fr1 = p_fr1.add_run('NGƯỜI PHÊ DUYỆT\n\n\n\n\n')
        r_fr1.bold = True
        p_fr1.add_run('{{NGUOI_KY}}').bold = True

    # Make sure we add Kính gửi if it's missing (before CAN_CU)
    has_kinhgui = False
    for p in doc.paragraphs:
        if 'Kính gửi' in p.text or '{{KINH_GUI}}' in p.text:
            has_kinhgui = True
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if 'Kính gửi' in c.text or '{{KINH_GUI}}' in c.text:
                    has_kinhgui = True
                    
    if not has_kinhgui:
        # Just insert it before {{CAN_CU}}
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for i, p in enumerate(c.paragraphs):
                        if '{{CAN_CU}}' in p.text:
                            p.insert_paragraph_before('Kính gửi: {{KINH_GUI}}').runs[0].bold = True
                            has_kinhgui = True
                            break
                    if has_kinhgui: break
                if has_kinhgui: break
            if has_kinhgui: break

    doc.save(dst)
    print("Done")

if __name__ == '__main__':
    prepare()
