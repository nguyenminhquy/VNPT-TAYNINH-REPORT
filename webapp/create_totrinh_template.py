import sys
import os
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_template():
    doc = Document()
    
    # 1. Page Margins (Left 30mm, Right 15mm, Top 20mm, Bottom 20mm) & Paper Size (A4)
    sections = doc.sections
    for section in sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
    
    # 2. Base Style: Times New Roman, Size 14, No Italics
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.italic = False
    
    # Configure paragraph spacing for the whole document
    style_normal.paragraph_format.space_after = Pt(6)
    
    # 3. Header Table (2 Columns)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    # A4 width = 210mm. Left margin = 30mm, Right margin = 15mm. Usable width = 165mm.
    # We split it roughly 70mm and 95mm for example, or equal 82.5mm each
    table.columns[0].width = Mm(70)
    table.columns[1].width = Mm(95)
    
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.space_after = Pt(0)
    
    r_left1 = p_left.add_run('TẬP ĐOÀN BƯU CHÍNH VIỄN THÔNG VIỆT NAM\n')
    r_left1.font.size = Pt(12)
    r_left1.italic = False
    
    r_left2 = p_left.add_run('VNPT TÂY NINH')
    r_left2.bold = True
    r_left2.font.size = Pt(13)
    r_left2.italic = False
    
    p_left_2 = cell_left.add_paragraph()
    p_left_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left_2.paragraph_format.space_after = Pt(0)
    r_left3 = p_left_2.add_run('Số: {{SO_TO_TRINH}}')
    r_left3.font.size = Pt(13)
    r_left3.italic = False
    
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.space_after = Pt(0)
    
    r_right1 = p_right.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n')
    r_right1.bold = True
    r_right1.font.size = Pt(13)
    r_right1.italic = False
    
    r_right2 = p_right.add_run('Độc lập - Tự do - Hạnh phúc')
    r_right2.bold = True
    r_right2.font.size = Pt(14)
    r_right2.underline = True
    r_right2.italic = False
    
    p_right_2 = cell_right.add_paragraph()
    p_right_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right_2.paragraph_format.space_after = Pt(0)
    r_right3 = p_right_2.add_run('{{NGAY_THANG}}')
    r_right3.font.size = Pt(13)
    r_right3.italic = False # NO ITALIC allowed per instruction
    
    doc.add_paragraph() # Spacing
    
    # 4. Title
    p_title1 = doc.add_paragraph()
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_after = Pt(0)
    r_t1 = p_title1.add_run('TỜ TRÌNH')
    r_t1.bold = True
    r_t1.font.size = Pt(16)
    r_t1.italic = False
    
    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_after = Pt(12)
    r_t2 = p_title2.add_run('V/v: {{TIEU_DE}}')
    r_t2.bold = True
    r_t2.font.size = Pt(14)
    r_t2.italic = False # NOT ITALIC
    
    # 5. Kinh Gui
    p_kinhgui = doc.add_paragraph()
    p_kinhgui.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_kg1 = p_kinhgui.add_run('Kính gửi: ')
    r_kg1.bold = True
    r_kg1.font.size = Pt(14)
    r_kg1.italic = False
    r_kg2 = p_kinhgui.add_run('{{KINH_GUI}}')
    r_kg2.bold = True
    r_kg2.font.size = Pt(14)
    r_kg2.italic = False
    
    # 6. Can Cu (Justify alignment)
    p_cancu = doc.add_paragraph('{{CAN_CU}}')
    p_cancu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 7. Noi Dung (Justify alignment)
    p_noidung = doc.add_paragraph('{{NOI_DUNG}}')
    p_noidung.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 8. De Xuat (Justify alignment)
    p_dexuat = doc.add_paragraph('{{DE_XUAT}}')
    p_dexuat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph() # Spacing before footer
    
    # 9. Footer: Nhan & Ky
    table_footer = doc.add_table(rows=1, cols=3)
    table_footer.autofit = False
    # Use 3 columns: Left (Nơi nhận), Middle (Người lập), Right (Người phê duyệt)
    table_footer.columns[0].width = Mm(50)
    table_footer.columns[1].width = Mm(50)
    table_footer.columns[2].width = Mm(65)
    
    # Nơi nhận
    cell_f_left = table_footer.cell(0, 0)
    p_fl1 = cell_f_left.paragraphs[0]
    p_fl1.paragraph_format.space_after = Pt(0)
    r_fl1 = p_fl1.add_run('Nơi nhận:\n')
    r_fl1.bold = True
    r_fl1.font.size = Pt(12)
    r_fl1.italic = False
    r_fl2 = p_fl1.add_run('{{NOI_NHAN}}')
    r_fl2.font.size = Pt(11)
    r_fl2.italic = False
    
    # Người lập
    cell_f_mid = table_footer.cell(0, 1)
    p_fm1 = cell_f_mid.paragraphs[0]
    p_fm1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fm1.paragraph_format.space_after = Pt(0)
    r_fm1 = p_fm1.add_run('NGƯỜI LẬP\n\n\n\n\n')
    r_fm1.bold = True
    r_fm1.italic = False
    r_fm2 = p_fm1.add_run('{{NGUOI_LAP}}')
    r_fm2.bold = True
    r_fm2.italic = False
    
    # Người phê duyệt
    cell_f_right = table_footer.cell(0, 2)
    p_fr1 = cell_f_right.paragraphs[0]
    p_fr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fr1.paragraph_format.space_after = Pt(0)
    r_fr1 = p_fr1.add_run('NGƯỜI PHÊ DUYỆT\n\n\n\n\n')
    r_fr1.bold = True
    r_fr1.italic = False
    r_fr2 = p_fr1.add_run('{{NGUOI_KY}}')
    r_fr2.bold = True
    r_fr2.italic = False

    os.makedirs('templates', exist_ok=True)
    out_path = os.path.join('templates', 'ToTrinh_Template.docx')
    doc.save(out_path)
    print(f"Created template at {out_path}")

if __name__ == '__main__':
    create_template()
