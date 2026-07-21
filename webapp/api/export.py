import json
import os
import tempfile
import traceback
from http.server import BaseHTTPRequestHandler
import requests
import re
from datetime import date, datetime
import math
import copy
from typing import Any, Sequence, Iterable
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell


def clean(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y %H:%M:%S")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, bool):
        return "Có" if value else "Không"
    return str(value).strip()


def trim_number(value: float, decimals: int = 5) -> str:
    if math.isclose(value, round(value), abs_tol=1e-10):
        return str(int(round(value)))
    return f"{value:.{decimals}f}".rstrip("0").rstrip(".")


def format_cell(cell: Cell) -> str:
    value = cell.value
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return clean(value)
    if not isinstance(value, (int, float)) or isinstance(value, bool):
        return clean(value)

    number_format = str(cell.number_format or "General").split(";")[0]
    if "%" in number_format:
        decimals_match = re.search(r"\.([0#]+)%", number_format)
        decimals = len(decimals_match.group(1)) if decimals_match else 0
        return f"{value * 100:.{decimals}f}%"

    decimal_match = re.search(r"\.([0#]+)", number_format)
    decimals = len(decimal_match.group(1)) if decimal_match else None
    use_grouping = "," in number_format.split(".")[0]
    if decimals is not None:
        return f"{value:,.{decimals}f}" if use_grouping else f"{value:.{decimals}f}"
    if use_grouping:
        return f"{value:,.0f}"
    return trim_number(float(value))


def worksheet_matrix(
    sheet: Any,
    min_row: int,
    max_row: int,
    min_col: int,
    max_col: int,
) -> list[list[str]]:
    return [
        [format_cell(sheet.cell(row=row, column=column)) for column in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def raw_matrix(
    sheet: Any,
    min_row: int,
    max_row: int,
    min_col: int,
    max_col: int,
) -> list[list[Any]]:
    return [
        [sheet.cell(row=row, column=column).value for column in range(min_col, max_col + 1)]
        for row in range(min_row, max_row + 1)
    ]


def load_sources() -> dict[str, Any]:
    missing = [filename for filename in FILES.values() if not (DATA_DIR / filename).is_file()]
    if missing:
        raise FileNotFoundError("Thiếu tệp nguồn: " + ", ".join(missing))

    sources: dict[str, Any] = {}
    for key, filename in FILES.items():
        sources[key] = load_workbook(DATA_DIR / filename, data_only=True, read_only=True)
    return sources


def first_run_properties(paragraph: Paragraph) -> Any | None:
    for run in paragraph._p.xpath(".//w:r"):
        run_properties = run.find(qn("w:rPr"))
        if run_properties is not None:
            return copy.deepcopy(run_properties)
    return None


def append_text_run(paragraph: Paragraph, value: str, run_properties: Any | None) -> None:
    lines = clean(value).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    run = OxmlElement("w:r")
    if run_properties is not None:
        run.append(copy.deepcopy(run_properties))

    for index, line in enumerate(lines):
        if index:
            run.append(OxmlElement("w:br"))
        text = OxmlElement("w:t")
        if line.startswith(" ") or line.endswith(" "):
            text.set(qn("xml:space"), "preserve")
        text.text = line
        run.append(text)
    paragraph._p.append(run)


def replace_paragraph(paragraph: Paragraph, value: Any) -> None:
    run_properties = first_run_properties(paragraph)
    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)
    append_text_run(paragraph, clean(value), run_properties)


def replace_cell(cell: _Cell, value: Any) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        paragraph_element = OxmlElement("w:p")
        cell._tc.append(paragraph_element)
        paragraphs = [Paragraph(paragraph_element, cell)]

    first = paragraphs[0]
    replace_paragraph(first, value)
    for paragraph in paragraphs[1:]:
        cell._tc.remove(paragraph._p)


def write_table_matrix(
    table: Table,
    matrix: Sequence[Sequence[Any]],
    start_row: int = 0,
    start_col: int = 0,
) -> None:
    touched: set[Any] = set()
    for source_row_index, source_row in enumerate(matrix):
        target_row_index = start_row + source_row_index
        if target_row_index >= len(table.rows):
            break
        cells = table.rows[target_row_index].cells
        for source_col_index, value in enumerate(source_row):
            target_col_index = start_col + source_col_index
            if target_col_index >= len(cells):
                break
            cell = cells[target_col_index]
            identity = cell._tc
            if identity in touched:
                continue
            replace_cell(cell, value)
            touched.add(identity)


def resize_table_rows(table: Table, desired_rows: int) -> None:
    while len(table.rows) < desired_rows:
        source_row = table.rows[-1]._tr
        table._tbl.append(copy.deepcopy(source_row))
    while len(table.rows) > desired_rows:
        table._tbl.remove(table.rows[-1]._tr)


def as_number(value: Any) -> float:
    if value is None or value == "":
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace("%", "").replace(",", "").strip()
    try:
        return float(text)
    except ValueError:
        return 0.0


def percent(value: Any, decimals: int = 2, stored_as_percent: bool = False) -> str:
    number = as_number(value)
    if not stored_as_percent:
        number *= 100
    return f"{number:.{decimals}f}%"


def integer(value: Any, dash_zero: bool = False) -> str:
    number = as_number(value)
    if dash_zero and math.isclose(number, 0.0, abs_tol=1e-9):
        return "-"
    return f"{round(number):,}"


def decimal(value: Any, decimals: int = 2) -> str:
    return f"{as_number(value):.{decimals}f}"


def evaluate_target(value: Any, target: float = 99.0) -> str:
    return "Đạt" if as_number(value) >= target else "Không đạt"


def update_mbb_fbb_mytv(document: DocumentType, sources: dict[str, Any]) -> None:
    mbb = sources["mbb"]
    fbb = sources["fbb"]
    mytv = sources["mytv"]

    common = worksheet_matrix(mbb["Kết quả chung"], 4, 6, 1, 4)
    for row in common[1:]:
        row[2] = decimal(row[2])
        row[3] = decimal(row[3])
    write_table_matrix(document.tables[1], common)

    comparison = worksheet_matrix(mbb["So sánh các tỉnh"], 3, 7, 1, 3)
    write_table_matrix(document.tables[2], comparison)

    mbb_detail = worksheet_matrix(mbb["Kết quả chi tiết"], 4, 12, 1, 8)
    write_table_matrix(document.tables[3], mbb_detail, start_row=3)

    fbb_detail = worksheet_matrix(fbb["Thông tin chung"], 2, 17, 1, 8)
    write_table_matrix(document.tables[3], fbb_detail, start_row=13)

    mytv_rows = raw_matrix(mytv["Sheet1"], 3, 16, 1, 8)
    mytv_detail: list[list[str]] = []
    for row in mytv_rows:
        total = row[6]
        mytv_detail.append(
            [
                clean(row[0]),
                clean(row[1]),
                clean(row[3]),
                clean(row[4]),
                clean(row[5]),
                clean(total),
                evaluate_target(total),
                clean(row[7]),
            ]
        )
    write_table_matrix(document.tables[3], mytv_detail, start_row=30)

    qos_explanation = worksheet_matrix(mbb["Giải trình QoS"], 4, 10, 1, 4)
    qoe_explanation = worksheet_matrix(mbb["Giải trình QoE"], 4, 8, 1, 4)
    plan_sheet_name = next(name for name in mbb.sheetnames if name.startswith("Dự kiến tuần"))
    plan = worksheet_matrix(mbb[plan_sheet_name], 3, 9, 1, 4)
    feedback = worksheet_matrix(mbb["Phản ánh khách hàng (PAKH)"], 4, 10, 1, 3)
    write_table_matrix(document.tables[4], qos_explanation)
    write_table_matrix(document.tables[5], qoe_explanation)
    write_table_matrix(document.tables[6], plan)
    write_table_matrix(document.tables[7], feedback)

    qos_sheet = fbb["Chi tiết QoS FBB"]
    write_table_matrix(document.tables[8], worksheet_matrix(qos_sheet, 1, 2, 1, 2))
    write_table_matrix(document.tables[9], worksheet_matrix(qos_sheet, 5, 6, 1, 7))
    write_table_matrix(document.tables[10], worksheet_matrix(qos_sheet, 9, 17, 1, 4))
    write_table_matrix(document.tables[11], worksheet_matrix(qos_sheet, 20, 43, 1, 5))
    write_table_matrix(
        document.tables[12],
        worksheet_matrix(fbb["Suy hao thuê bao"], 2, 25, 1, 7),
    )

    plan_week_match = re.search(r"(\d+)$", plan_sheet_name)
    if plan_week_match:
        replace_paragraph(
            document.paragraphs[18],
            f"Công việc dự kiến tuần {plan_week_match.group(1)}:",
        )

    feedback_cutoff = next(
        (clean(row[1]) for row in feedback[1:] if "đến" in clean(row[1]).lower()),
        "",
    )
    cutoff_match = re.search(r"(\d{1,2}/\d{1,2}/\d{4})", feedback_cutoff)
    if cutoff_match:
        cutoff = datetime.strptime(cutoff_match.group(1), "%d/%m/%Y")
        replace_paragraph(
            document.paragraphs[20],
            f"Thời gian lấy báo cáo: 01/{cutoff.month:02d}/{cutoff.year} – {cutoff:%d/%m/%Y}",
        )


def mll_table_matrix(sheet: Any) -> tuple[list[list[str]], dict[str, Any]]:
    raw = raw_matrix(sheet, 2, 12, 1, 18)
    title = clean(raw[0][0])
    matrix: list[list[str]] = [[title]]

    source_columns = [0, 1, 2, 3, 4, 5, 7, 8, 9, 11, 12, 13, 15, 16, 17]
    for row in raw[3:11]:
        target: list[str] = []
        for target_index, source_index in enumerate(source_columns):
            value = row[source_index]
            if target_index in {0, 1}:
                target.append(clean(value))
            elif target_index == 2:
                target.append(integer(value))
            elif 3 <= target_index <= 11:
                target.append(integer(value, dash_zero=True))
            elif target_index in {12, 13}:
                target.append(integer(value))
            else:
                target.append(decimal(value))
        matrix.append(target)

    overall = raw[3]
    teams = raw[4:11]
    metrics = {
        "title": title,
        "week": re.search(r"TUẦN\s+(\d+)", title, re.IGNORECASE),
        "total": as_number(overall[15]),
        "average": as_number(overall[17]),
        "teams": [(clean(row[1]), as_number(row[17])) for row in teams],
        "cause_power": as_number(overall[3]) + as_number(overall[7]) + as_number(overall[11]),
        "cause_equipment": as_number(overall[4]) + as_number(overall[8]) + as_number(overall[12]),
        "cause_transmission": as_number(overall[5]) + as_number(overall[9]) + as_number(overall[13]),
    }
    return matrix, metrics


def update_mll(document: DocumentType, sources: dict[str, Any]) -> str:
    matrix, metrics = mll_table_matrix(sources["mll"]["BC MLL tuần"])
    replace_cell(document.tables[13].rows[0].cells[0], matrix[0][0])
    write_table_matrix(document.tables[13], matrix[1:], start_row=3)

    week = metrics["week"].group(1) if metrics["week"] else ""
    replace_paragraph(document.paragraphs[36], f"Tổng thời gian mất liên lạc: {metrics['total']:,.0f} phút.")
    replace_paragraph(document.paragraphs[37], f"MLL trung bình/1 BTS: {metrics['average']:.2f} phút.")

    if week:
        replace_paragraph(document.paragraphs[41], f"Đánh giá thời gian mất liên lạc vô tuyến tuần {week}:")
        replace_paragraph(
            document.paragraphs[50],
            f"Nguyên nhân chi tiết các trạm MLL trong tuần {week} năm 2026 và các đánh giá, giải pháp khắc phục (Theo phụ lục 01 đính kèm)",
        )
        replace_paragraph(
            document.paragraphs[121],
            f"GIẢI TRÌNH NGUYÊN NHÂN MẤT LIÊN LẠC TRẠM TUẦN {week}",
        )

    achieved = sum(1 for _, average in metrics["teams"] if average <= 3.0)
    replace_paragraph(
        document.paragraphs[42],
        f"{achieved}/7 THT có thời gian mất liên lạc đáp ứng chỉ tiêu của VTT (≤3 phút).",
    )
    highest = sorted(metrics["teams"], key=lambda item: item[1], reverse=True)[:3]
    highest_text = ", ".join(f"THT {name} ({value:.2f} phút/1 trạm)" for name, value in highest)
    replace_paragraph(
        document.paragraphs[43],
        "Thời gian mất liên lạc trung bình trên 1 trạm BTS cao nhất: " + highest_text + ".",
    )

    total = metrics["total"] or 1.0
    replace_paragraph(document.paragraphs[45], f"MLL do lỗi nguồn ({metrics['cause_power'] / total:.0%})")
    replace_paragraph(document.paragraphs[46], f"MLL do lỗi thiết bị ({metrics['cause_equipment'] / total:.0%})")
    replace_paragraph(document.paragraphs[47], f"MLL do lỗi truyền dẫn ({metrics['cause_transmission'] / total:.0%})")
    return week


def update_ispeed(document: DocumentType, sources: dict[str, Any]) -> None:
    sheet = sources["ispeed"]["Báo cáo"]
    raw = raw_matrix(sheet, 1, 9, 1, 11)
    matrix: list[list[str]] = []
    for row_index, row in enumerate(raw):
        if row_index == 0:
            matrix.append([clean(value) for value in row])
            continue
        matrix.append(
            [
                clean(row[0]),
                clean(row[1]),
                integer(row[2]),
                integer(row[3]),
                integer(row[4]),
                percent(row[5]),
                integer(row[6]),
                integer(row[7]),
                percent(row[8]),
                integer(row[9]),
                percent(row[10]),
            ]
        )
    write_table_matrix(document.tables[14], matrix)

    report_date = clean(sheet.cell(row=12, column=2).value)
    if report_date:
        replace_paragraph(document.paragraphs[55], "Thời gian lấy báo cáo: " + report_date)

    total = raw[-1]
    replace_paragraph(
        document.paragraphs[59],
        f"Công tác đo kiểm i-Speed đã thực hiện {integer(total[4])}/{integer(total[3])} mẫu, đạt {percent(total[5])}/Tháng kế hoạch.",
    )
    replace_paragraph(
        document.paragraphs[60],
        f"Công tác đo kiểm SpeedTest đã thực hiện {integer(total[7])}/{integer(total[6])} mẫu, đạt {percent(total[8])}/Tháng kế hoạch.",
    )
    replace_paragraph(
        document.paragraphs[61],
        f"Kết quả mẫu đo 5G SpeedTest đã thực hiện {integer(total[9])}/{integer(total[7])} mẫu, đạt {percent(total[10])}/Tổng mẫu đã đo.",
    )


def format_5s_matrix(rows: list[list[Any]]) -> list[list[str]]:
    result: list[list[str]] = []
    for row_index, row in enumerate(rows):
        if row_index <= 1:
            result.append([clean(value) for value in row])
            continue
        result.append(
            [
                clean(row[0]),
                clean(row[1]),
                integer(row[2]),
                integer(row[3]),
                integer(row[4]),
                percent(row[5]),
            ]
        )
    return result


def update_5s(document: DocumentType, sources: dict[str, Any]) -> None:
    sheet = sources["5s"]["Sheet1"]
    station = format_5s_matrix(raw_matrix(sheet, 1, 10, 1, 6))
    air_conditioning = format_5s_matrix(raw_matrix(sheet, 14, 23, 1, 6))
    ap_otb = format_5s_matrix(raw_matrix(sheet, 26, 35, 1, 6))
    write_table_matrix(document.tables[15], station)
    write_table_matrix(document.tables[16], ap_otb)
    write_table_matrix(document.tables[17], air_conditioning)

    source_date = datetime.fromtimestamp((DATA_DIR / FILES["5s"]).stat().st_mtime)
    value = f"Thời gian lấy báo cáo: {source_date:%d/%m/%Y}"
    for index in (64, 70, 76):
        replace_paragraph(document.paragraphs[index], value)


def parse_xlsc_title(title: str) -> tuple[str, str, int, int] | None:
    match = re.search(
        r"\((\d{2})-(\d{2})-(\d{4})\s*-\s*(\d{2})-(\d{2})-(\d{4})\)",
        title,
    )
    if not match:
        return None
    start = f"{match.group(1)}/{match.group(2)}/{match.group(3)}"
    end = f"{match.group(4)}/{match.group(5)}/{match.group(6)}"
    return start, end, int(match.group(5)), int(match.group(6))


def xlsc_matrix(sheet: Any) -> list[list[str]]:
    raw = raw_matrix(sheet, 1, 10, 1, 10)
    result: list[list[str]] = []
    for row_index, row in enumerate(raw):
        if row_index <= 1:
            result.append([clean(value) for value in row])
            continue
        result.append(
            [
                clean(row[0]),
                integer(row[1]),
                integer(row[2]),
                integer(row[3]),
                integer(row[4]),
                percent(row[5], stored_as_percent=True),
                integer(row[6]),
                integer(row[7]),
                integer(row[8]),
                percent(row[9], stored_as_percent=True),
            ]
        )
    return result


def update_xlsc_summary(
    document: DocumentType,
    paragraph_start: int,
    total_row: Sequence[Any],
    date_range: tuple[str, str, int, int] | None,
) -> None:
    if date_range:
        replace_paragraph(
            document.paragraphs[paragraph_start],
            f"Kỳ báo cáo: {date_range[0]} – {date_range[1]}",
        )
    replace_paragraph(document.paragraphs[paragraph_start + 1], f"Tổng phiếu giao: {integer(total_row[1])} phiếu")
    replace_paragraph(
        document.paragraphs[paragraph_start + 2],
        f"Hoàn thành: {integer(total_row[2])}/{integer(total_row[1])} phiếu",
    )
    replace_paragraph(document.paragraphs[paragraph_start + 3], f"Hoàn thành đúng hạn: {integer(total_row[3])} phiếu")
    replace_paragraph(document.paragraphs[paragraph_start + 4], f"Hoàn thành quá hạn: {integer(total_row[4])} phiếu")
    replace_paragraph(
        document.paragraphs[paragraph_start + 5],
        f"Tỉ lệ đúng hạn: {percent(total_row[5], stored_as_percent=True)}",
    )
    replace_paragraph(document.paragraphs[paragraph_start + 6], f"Phiếu tồn quá hạn: {integer(total_row[8])} phiếu")


def update_xlsc(document: DocumentType, sources: dict[str, Any]) -> None:
    workbook = sources["xlsc"]
    mappings = [
        ("XLSC MANE", 18, 81),
        ("XLSC ACCESS", 19, 94),
        ("XLSC VÔ TUYẾN", 20, 105),
    ]
    report_month: tuple[int, int] | None = None
    for sheet_name, table_index, paragraph_start in mappings:
        sheet = workbook[sheet_name]
        matrix = xlsc_matrix(sheet)
        write_table_matrix(document.tables[table_index], matrix)
        date_range = parse_xlsc_title(clean(sheet.cell(row=1, column=1).value))
        if date_range:
            report_month = (date_range[2], date_range[3])
        total_row = raw_matrix(sheet, 10, 10, 1, 10)[0]
        update_xlsc_summary(document, paragraph_start, total_row, date_range)

    if report_month:
        replace_paragraph(
            document.paragraphs[79],
            f"KẾT QUẢ THỰC HIỆN PHIẾU SỰ CỐ CHUYÊN ĐỀ 5 THÁNG {report_month[0]} NĂM {report_month[1]}:",
        )


def update_appendix(document: DocumentType, sources: dict[str, Any]) -> None:
    sheet = sources["appendix"]["Báo Cáo Sự Cố Trạm"]
    data_rows = [
        row
        for row in range(5, sheet.max_row + 1)
        if str(sheet.cell(row=row, column=1).value or "").strip().isdigit()
    ]
    last_row = max(data_rows, default=4)
    matrix = worksheet_matrix(sheet, 4, last_row, 1, 10)
    table = document.tables[22]
    resize_table_rows(table, len(matrix))
    write_table_matrix(table, matrix)


def iter_story_roots(document: DocumentType) -> Iterable[Any]:
    yield document.element
    seen: set[int] = set()
    for section in document.sections:
        for story in (section.header, section.footer):
            identity = id(story.part)
            if identity not in seen:
                seen.add(identity)
                yield story._element


def flatten_simple_link_fields(root: Any) -> int:
    count = 0
    instruction_attribute = qn("w:instr")
    for field in list(root.iter(qn("w:fldSimple"))):
        if not clean(field.get(instruction_attribute)).upper().startswith("LINK "):
            continue
        parent = field.getparent()
        insert_at = parent.index(field)
        for child in list(field):
            field.remove(child)
            parent.insert(insert_at, child)
            insert_at += 1
        parent.remove(field)
        count += 1
    return count


def flatten_complex_link_fields(root: Any) -> int:
    contexts: list[dict[str, Any]] = []
    completed: list[dict[str, Any]] = []

    for run in list(root.iter(qn("w:r"))):
        for child in list(run):
            if child.tag == qn("w:fldChar"):
                marker = child.get(qn("w:fldCharType"))
                if marker == "begin":
                    contexts.append({"instruction": [], "code_runs": {run}, "phase": "code"})
                elif marker == "separate" and contexts:
                    contexts[-1]["code_runs"].add(run)
                    contexts[-1]["phase"] = "result"
                elif marker == "end" and contexts:
                    context = contexts.pop()
                    context["end_run"] = run
                    completed.append(context)
            elif child.tag == qn("w:instrText") and contexts:
                context = contexts[-1]
                if context["phase"] == "code":
                    context["instruction"].append(child.text or "")
                    context["code_runs"].add(run)

    # The retained template contains one unterminated 5S LINK field. Its code
    # can still be safely removed because the linked result table follows it.
    completed.extend(contexts)

    link_contexts = [
        context
        for context in completed
        if "".join(context["instruction"]).lstrip().upper().startswith("LINK ")
    ]
    runs_to_clean: set[Any] = set()
    for context in link_contexts:
        runs_to_clean.update(context["code_runs"])
        if context.get("end_run") is not None:
            runs_to_clean.add(context["end_run"])

    meaningful_tags = {
        qn("w:t"),
        qn("w:tab"),
        qn("w:br"),
        qn("w:drawing"),
        qn("w:object"),
        qn("w:pict"),
    }
    for run in runs_to_clean:
        for node in list(run.iter()):
            if node.tag not in {qn("w:fldChar"), qn("w:instrText")}:
                continue
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
        if not any(node.tag in meaningful_tags for node in run.iter()):
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)

    return len(link_contexts)


def flatten_link_fields(document: DocumentType) -> int:
    count = 0
    for root in iter_story_roots(document):
        count += flatten_simple_link_fields(root)
        count += flatten_complex_link_fields(root)
    return count


def replace_report_week(document: DocumentType, week: str) -> None:
    if not week:
        return
    plan_match = re.search(r"tuần\s+(\d+)", document.paragraphs[18].text, re.IGNORECASE)
    plan_week = plan_match.group(1) if plan_match else str(int(week) + 1)
    replacements = {
        1: f"V/v thực hiện công việc trọng tâm trong tuần {week} năm 2026",
        2: f"và kế hoạch thực hiện nhiệm vụ tuần {plan_week}",
        4: f"Trung tâm Hạ tầng báo cáo kết quả thực hiện công việc trọng tâm trong tuần {week} năm 2026 như sau:",
        116: f"Trên đây là báo cáo kết quả thực hiện công việc tuần {week} năm 2026.",
    }
    for index, value in replacements.items():
        replace_paragraph(document.paragraphs[index], value)


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get('Content-Length', 0))
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data)
            
            secret = data.get("secret")
            if secret != os.environ.get("PYTHON_API_SECRET", "vnpt-secret-key"):
                self.send_response(401)
                self.end_headers()
                self.wfile.write(b"Unauthorized")
                return
                
            blob_urls = data.get("blob_urls", {})
            
            with tempfile.TemporaryDirectory() as tmpdir:
                sources = {}
                for key, url in blob_urls.items():
                    resp = requests.get(url)
                    resp.raise_for_status()
                    filepath = os.path.join(tmpdir, f"{key}.xlsx")
                    with open(filepath, "wb") as f:
                        f.write(resp.content)
                    sources[key] = load_workbook(filepath, data_only=True, read_only=True)
                
                template_path = os.path.join(os.path.dirname(__file__), "templates", "template.docx")
                document = Document(template_path)
                
                update_mbb_fbb_mytv(document, sources)
                week = update_mll(document, sources)
                update_ispeed(document, sources)
                update_5s(document, sources)
                update_xlsc(document, sources)
                update_appendix(document, sources)
                replace_report_week(document, week)
                flatten_link_fields(document)
                
                out_path = os.path.join(tmpdir, "output.docx")
                document.save(out_path)
                
                with open(out_path, "rb") as f:
                    docx_bytes = f.read()
                    
                self.send_response(200)
                self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.end_headers()
                self.wfile.write(docx_bytes)
                
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-type', 'text/plain')
            self.end_headers()
            self.wfile.write(str(traceback.format_exc()).encode('utf-8'))
