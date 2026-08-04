import docx
import sys
import os

doc_path = r'c:\Users\quyng\Documents\TOOL TỔNG HỢP BÁO CÁO\VNPT REPORT\webapp\templates\TOTRINH\11_Mau_Giay_trieu_tap.docx'
doc = docx.Document(doc_path)

# Find the paragraph starting with "Giám đốc …….quyết định triệu tập"
# and remove it and subsequent ones until "Đề nghị các đơn vị phối hợp triển khai để Hội nghị đạt kết quả./."
start_idx = -1
end_idx = -1

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('Giám đốc'):
        start_idx = i
    if text.startswith('Đề nghị các đơn vị'):
        end_idx = i
        break

if start_idx != -1 and end_idx != -1:
    # Delete paragraphs from end_idx down to start_idx
    for i in range(end_idx, start_idx - 1, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)

    # Insert a new placeholder paragraph at start_idx
    new_p = doc.paragraphs[start_idx-1].insert_paragraph_before('(4)')
    new_p.style = doc.paragraphs[start_idx-1].style

doc.save(doc_path)
print('Modified Mẫu 11 template successfully.')
