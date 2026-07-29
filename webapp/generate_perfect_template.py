import sys
import os
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_template():
    doc = Document()
    
    # 1. Page Margins (Left 30mm, Right 15mm, Top 20mm, Bottom 20mm) & Paper Size (A4)
    sections = doc.sections
    for section in sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
    
    # 2. Base Style: Times New Roman, Size 14, No Italics
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.italic = False
    
    style_normal.paragraph_format.space_after = Pt(6)
    
    # 3. Header Table (2 Columns)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    table.columns[0].width = Mm(75)
    table.columns[1].width = Mm(90)
    
    # Left Header
    cell_left = table.cell(0, 0)
    
    p_l1 = cell_left.paragraphs[0]
    p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l1.paragraph_format.space_after = Pt(0)
    r_l1 = p_l1.add_run('TẬP ĐOÀN BƯU CHÍNH VIỄN THÔNG VIỆT NAM')
    r_l1.font.size = Pt(12)
    
    p_l2 = cell_left.add_paragraph()
    p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l2.paragraph_format.space_after = Pt(0)
    r_l2 = p_l2.add_run('VNPT TÂY NINH')
    r_l2.bold = True
    r_l2.font.size = Pt(13)
    
    p_l3 = cell_left.add_paragraph()
    p_l3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_l3.paragraph_format.space_after = Pt(0)
    r_l3 = p_l3.add_run('Số: {{SO_TO_TRINH}}')
    r_l3.font.size = Pt(13)
    
    # Right Header
    cell_right = table.cell(0, 1)
    
    p_r1 = cell_right.paragraphs[0]
    p_r1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r1.paragraph_format.space_after = Pt(0)
    r_r1 = p_r1.add_run('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM')
    r_r1.bold = True
    r_r1.font.size = Pt(13)
    
    p_r2 = cell_right.add_paragraph()
    p_r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r2.paragraph_format.space_after = Pt(0)
    r_r2 = p_r2.add_run('Độc lập - Tự do - Hạnh phúc')
    r_r2.bold = True
    r_r2.font.size = Pt(14)
    r_r2.underline = True
    
    p_r3 = cell_right.add_paragraph()
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_r3.paragraph_format.space_after = Pt(0)
    r_r3 = p_r3.add_run('{{NGAY_THANG}}')
    r_r3.font.size = Pt(13)
    r_r3.italic = False
    
    doc.add_paragraph() # Spacing
    
    # 4. Title
    p_title1 = doc.add_paragraph()
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_after = Pt(0)
    r_t1 = p_title1.add_run('TỜ TRÌNH')
    r_t1.bold = True
    r_t1.font.size = Pt(16)
    
    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_after = Pt(12)
    r_t2 = p_title2.add_run('V/v: {{TIEU_DE}}')
    r_t2.bold = True
    r_t2.font.size = Pt(14)
    
    # 5. Kinh Gui
    p_kinhgui = doc.add_paragraph()
    p_kinhgui.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_kg1 = p_kinhgui.add_run('Kính gửi: ')
    r_kg1.bold = True
    r_kg1.font.size = Pt(14)
    r_kg2 = p_kinhgui.add_run('{{KINH_GUI}}')
    r_kg2.bold = True
    r_kg2.font.size = Pt(14)
    
    # 6. Can Cu
    p_cancu = doc.add_paragraph('{{CAN_CU}}')
    p_cancu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 7. Noi Dung
    p_noidung = doc.add_paragraph('{{NOI_DUNG}}')
    p_noidung.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 8. De Xuat
    p_dexuat = doc.add_paragraph('{{DE_XUAT}}')
    p_dexuat.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph() # Spacing before footer
    
    # 9. Footer: Nhan & Ky
    table_footer = doc.add_table(rows=1, cols=3)
    table_footer.autofit = False
    table_footer.columns[0].width = Mm(50)
    table_footer.columns[1].width = Mm(50)
    table_footer.columns[2].width = Mm(65)
    
    # Nơi nhận
    cell_f_left = table_footer.cell(0, 0)
    p_fl1 = cell_f_left.paragraphs[0]
    p_fl1.paragraph_format.space_after = Pt(0)
    r_fl1 = p_fl1.add_run('Nơi nhận:')
    r_fl1.bold = True
    r_fl1.font.size = Pt(12)
    p_fl2 = cell_f_left.add_paragraph()
    p_fl2.paragraph_format.space_after = Pt(0)
    r_fl2 = p_fl2.add_run('{{NOI_NHAN}}')
    r_fl2.font.size = Pt(11)
    
    # Người lập
    cell_f_mid = table_footer.cell(0, 1)
    p_fm1 = cell_f_mid.paragraphs[0]
    p_fm1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fm1.paragraph_format.space_after = Pt(0)
    r_fm1 = p_fm1.add_run('NGƯỜI LẬP')
    r_fm1.bold = True
    
    # Spacing for signature
    cell_f_mid.add_paragraph().paragraph_format.space_after = Pt(0)
    cell_f_mid.add_paragraph().paragraph_format.space_after = Pt(0)
    cell_f_mid.add_paragraph().paragraph_format.space_after = Pt(0)
    cell_f_mid.add_paragraph().paragraph_format.space_after = Pt(0)
    
    p_fm2 = cell_f_mid.add_paragraph()
    p_fm2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fm2.paragraph_format.space_after = Pt(0)
    r_fm2 = p_fm2.add_run('{{NGUOI_LAP}}')
    r_fm2.bold = True
    
    # Người phê duyệt
    cell_f_right = table_footer.cell(0, 2)
    p_fr1 = cell_f_right.paragraphs[0]
    p_fr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fr1.paragraph_format.space_after = Pt(0)
    r_fr1 = p_fr1.add_run('NGƯỜI PHÊ DUYỆT')
    r_fr1.bold = True
    
    # Spacing for signature
    cell_f_right.add_paragraph().paragraph_format.space_after = Pt(0)
    cell_f_right.add_paragraph().paragraph_format.space_after = Pt(0)
    cell_f_right.add_paragraph().paragraph_format.space_after = Pt(0)
    cell_f_right.add_paragraph().paragraph_format.space_after = Pt(0)
    
    p_fr2 = cell_f_right.add_paragraph()
    p_fr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fr2.paragraph_format.space_after = Pt(0)
    r_fr2 = p_fr2.add_run('{{NGUOI_KY}}')
    r_fr2.bold = True

    # FORCE removal of all italics
    for p in doc.paragraphs:
        for r in p.runs:
            r.italic = False
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.italic = False

    os.makedirs('templates', exist_ok=True)
    out_path = os.path.join('templates', 'ToTrinh_Template.docx')
    doc.save(out_path)
    print(f"Created flawless template at {out_path}")

if __name__ == '__main__':
    create_template()
